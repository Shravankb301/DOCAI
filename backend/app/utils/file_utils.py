from fastapi import UploadFile
import os
import mimetypes
from typing import List, Set, Dict, Any, Tuple, Optional
import importlib.util
import chardet
import io
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Define allowed file extensions and their corresponding MIME types
ALLOWED_EXTENSIONS: Set[str] = {
    '.txt', '.pdf', '.doc', '.docx', '.rtf', '.md', '.csv', '.json'
}

# Define MIME types for binary files that need special handling
BINARY_MIME_TYPES: Set[str] = {
    'application/pdf',
    'application/msword',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
}

# Maximum file size (10 MB)
MAX_FILE_SIZE: int = 10 * 1024 * 1024  # 10 MB in bytes

# Check if optional dependencies are available
PYPDF2_AVAILABLE = importlib.util.find_spec("PyPDF2") is not None
DOCX_AVAILABLE = importlib.util.find_spec("docx") is not None
PANDAS_AVAILABLE = importlib.util.find_spec("pandas") is not None

def validate_file(file: UploadFile) -> Tuple[bool, str]:
    """
    Validate if the uploaded file meets the requirements.
    
    Args:
        file: The uploaded file to validate
        
    Returns:
        Tuple[bool, str]: (True if file is valid, error message if invalid)
    """
    # Check file extension
    _, file_extension = os.path.splitext(file.filename)
    if file_extension.lower() not in ALLOWED_EXTENSIONS:
        return False, f"File type {file_extension} not allowed. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"
    
    # Check file size (this requires reading the file)
    # Note: This approach loads the file into memory, which may not be ideal for very large files
    file.file.seek(0, os.SEEK_END)
    file_size = file.file.tell()
    file.file.seek(0)  # Reset file pointer
    
    if file_size > MAX_FILE_SIZE:
        return False, f"File size exceeds maximum allowed size of {MAX_FILE_SIZE/1024/1024} MB"
    
    # Check if file is empty
    if file_size == 0:
        return False, "File is empty"
    
    return True, ""

def detect_file_type(file_path: str) -> Dict[str, Any]:
    """
    Detect file type and encoding.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Dict: Information about the file type and encoding
    """
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()
    
    # Get MIME type
    mime_type, _ = mimetypes.guess_type(file_path)
    
    # Detect if file is binary or text
    is_binary = False
    encoding = 'utf-8'
    
    # Check if it's a known binary type
    if mime_type in BINARY_MIME_TYPES:
        is_binary = True
    else:
        # Try to detect encoding for potentially text files
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read(4096)  # Read first 4KB to detect encoding
                result = chardet.detect(raw_data)
                encoding = result['encoding'] or 'utf-8'
                confidence = result['confidence']
                
                # If confidence is low or encoding is binary-like, treat as binary
                if confidence < 0.7 or encoding.lower() in ['binary', 'unknown-8bit', None]:
                    is_binary = True
        except Exception as e:
            logger.error(f"Error detecting file encoding: {str(e)}")
            is_binary = True  # Default to binary if detection fails
    
    return {
        'file_extension': file_extension,
        'mime_type': mime_type,
        'is_binary': is_binary,
        'encoding': encoding
    }

def get_file_content(file_path: str) -> str:
    """
    Extract text content from a file.
    
    Args:
        file_path: Path to the file
        
    Returns:
        str: Text content of the file
    """
    try:
        # Detect file type and encoding
        file_info = detect_file_type(file_path)
        file_extension = file_info['file_extension']
        is_binary = file_info['is_binary']
        encoding = file_info['encoding']
        
        # Handle different file types
        if file_extension == '.pdf' and PYPDF2_AVAILABLE:
            return extract_text_from_pdf(file_path)
        elif file_extension == '.docx' and DOCX_AVAILABLE:
            return extract_text_from_docx(file_path)
        elif file_extension == '.csv' and PANDAS_AVAILABLE:
            return extract_text_from_csv(file_path)
        elif file_extension == '.json':
            return extract_text_from_json(file_path)
        elif is_binary:
            logger.warning(f"File {file_path} appears to be binary but no specific handler is available")
            return f"[Binary file detected: {os.path.basename(file_path)}. No suitable text extractor available.]"
        else:
            # Handle text files with detected encoding
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                return f.read()
    except Exception as e:
        logger.error(f"Error reading file {file_path}: {str(e)}")
        return f"[Error reading file: {str(e)}]"

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file using PyPDF2.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        str: Extracted text content
    """
    try:
        import PyPDF2
        
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Check if PDF is encrypted
            if pdf_reader.is_encrypted:
                return "[PDF is encrypted and cannot be read without a password]"
                
            for page_num in range(len(pdf_reader.pages)):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text() or "[No text could be extracted from this page]"
                    text += f"--- Page {page_num + 1} ---\n{page_text}\n\n"
                except Exception as e:
                    text += f"[Error extracting text from page {page_num + 1}: {str(e)}]\n\n"
        
        return text if text.strip() else "[No readable text found in PDF]"
    except ImportError:
        logger.warning("PyPDF2 is not installed. Cannot extract text from PDF.")
        return "[PyPDF2 is not installed. Cannot extract text from PDF.]"
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
        return f"[Error extracting text from PDF: {str(e)}]"

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file using python-docx.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        str: Extracted text content
    """
    try:
        from docx import Document
        
        document = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for para in document.paragraphs:
            text += para.text + "\n"
        
        # Extract text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text if text.strip() else "[No readable text found in DOCX]"
    except ImportError:
        logger.warning("python-docx is not installed. Cannot extract text from DOCX.")
        return "[python-docx is not installed. Cannot extract text from DOCX.]"
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
        return f"[Error extracting text from DOCX: {str(e)}]"

def extract_text_from_csv(file_path: str) -> str:
    """
    Extract text from a CSV file using pandas.
    
    Args:
        file_path: Path to the CSV file
        
    Returns:
        str: Extracted text content
    """
    try:
        import pandas as pd
        
        # Try to detect encoding
        file_info = detect_file_type(file_path)
        encoding = file_info['encoding']
        
        # Read CSV file
        df = pd.read_csv(file_path, encoding=encoding, error_bad_lines=False, warn_bad_lines=True)
        
        # Convert to string representation
        text = "CSV Data:\n"
        text += df.to_string(index=False, max_rows=100)
        
        if len(df) > 100:
            text += f"\n\n[Note: Only showing first 100 rows out of {len(df)} total rows]"
            
        return text
    except ImportError:
        logger.warning("pandas is not installed. Cannot extract text from CSV.")
        return "[pandas is not installed. Cannot extract text from CSV.]"
    except Exception as e:
        logger.error(f"Error extracting text from CSV {file_path}: {str(e)}")
        return f"[Error extracting text from CSV: {str(e)}]"

def extract_text_from_json(file_path: str) -> str:
    """
    Extract text from a JSON file.
    
    Args:
        file_path: Path to the JSON file
        
    Returns:
        str: Extracted text content
    """
    try:
        import json
        
        # Try to detect encoding
        file_info = detect_file_type(file_path)
        encoding = file_info['encoding']
        
        with open(file_path, 'r', encoding=encoding, errors='replace') as f:
            data = json.load(f)
        
        # Convert to formatted string
        return json.dumps(data, indent=2)
    except Exception as e:
        logger.error(f"Error extracting text from JSON {file_path}: {str(e)}")
        return f"[Error extracting text from JSON: {str(e)}]" 